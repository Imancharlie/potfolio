# projects/views.py
from django.shortcuts import render, get_object_or_404, redirect
from django.http import FileResponse
from django.shortcuts import render, redirect
from .models import Project, Feedback
from django.contrib import messages
from django.http import JsonResponse

from django.http import HttpResponse
from docx import Document
from potfolio_app.models import Project  # Replace with your model
from django.http import HttpResponse
from docx import Document
from .models import Project  # Adjust the import based on your app's name

def export_to_word(request):
    # Fetch all projects from the database
    projects = Project.objects.all()

    # Create a new Word document
    doc = Document()
    doc.add_heading('Projects Export', 0)

    # Loop through each project and add it to the document
    for project in projects:
        doc.add_paragraph(f"Title: {project.title}")
        doc.add_paragraph(f"Slug: {project.slug}")
        doc.add_paragraph(f"Description: {project.description}")
        doc.add_paragraph(f"Short Description: {project.short_description}")
        doc.add_paragraph(f"Created At: {project.created_at}")
        doc.add_paragraph(f"Updated At: {project.updated_at}")
        doc.add_paragraph(f"Has Downloadable File: {'Yes' if project.has_download else 'No'}")
        if project.download_file:
            doc.add_paragraph(f"Download File: {project.download_file.url}")
        if project.external_url:
            doc.add_paragraph(f"External URL: {project.external_url}")
        doc.add_paragraph()  # Add a blank line between each project

    # Create an HTTP response with the Word document content
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="projects_export.docx"'

    # Save the document to the response and return it
    doc.save(response)
    return response


def home(request):
    projects = Project.objects.all().order_by('-created_at')

    if request.method == "POST":
        name = request.POST.get('name')
        email = request.POST.get('email')
        subject = request.POST.get('subject')
        message = request.POST.get('message')

        if name and email and subject and message:
            Feedback.objects.create(name=name, email=email, subject=subject, message=message)
            return redirect('home')

    return render(request, 'projects/home.html', {'projects': projects})

def project_detail(request, slug):
    project = get_object_or_404(Project, slug=slug)
        
    return render(request, 'projects/project_detail.html', {'project': project})

def download_file(request, slug):
    project = get_object_or_404(Project, slug=slug)
    if project.download_file:
        return FileResponse(project.download_file.open(), as_attachment=True)
    return redirect('project_detail', slug=slug)